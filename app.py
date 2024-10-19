import streamlit as st
import pandas as pd
import json
import uuid  # For generating unique IDs
import os
from sentence_transformers import SentenceTransformer
import chromadb
import google.generativeai as genai
from IPython.display import Markdown
from chunking import RecursiveTokenChunker, LLMAgenticChunkerv2, ProtonxSemanticChunker
from utils import process_batch, divide_dataframe
from search import vector_search, keywords_search, hyde_search
from llms.localLllms import run_ollama_container, run_ollama_model, OLLAMA_MODEL_OPTIONS, GGUF_MODEL_OPTIONS
from llms.onlinellms import OnlineLLMs
import time
import pdfplumber  # PDF extraction
import io
from docx import Document  # DOCX extraction
from components import notify
from constant import NO_CHUNKING, EN, VI, NONE, USER, ASSISTANT, ENGLISH, VIETNAMESE, ONLINE_LLM, LOCAL_LLM, GEMINI, DEFAULT_LOCAL_LLM, OPENAI, DB
from graph_rag import GraphRAG
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.schema import Document as langchainDocument
from langchain_openai import ChatOpenAI
from collection_management import list_collection
from dotenv import load_dotenv
load_dotenv()

def clear_session_state():
    for key in st.session_state.keys():
        del st.session_state[key]

# Button to clear all session state
# st.sidebar.header("Clear Session State")
# if st.sidebar.button("Clear Session State"):
#     clear_session_state()
#     st.success("Session state has been cleared successfully!")

llm_options = {
    "Online": "Online",
    "Local (Ollama)": "Local (Ollama)"
}


# Initialize the page
st.header("Drag and Drop RAG")
st.markdown("Design your own chatbot using the RAG system.")
st.logo("https://storage.googleapis.com/mle-courses-prod/users/61b6fa1ba83a7e37c8309756/private-files/678dadd0-603b-11ef-b0a7-998b84b38d43-ProtonX_logo_horizontally__1_.png")


# --- Initialize session state for language choice and model embedding
if "language" not in st.session_state:
    st.session_state.language = None  # Default language is English
if "embedding_model" not in st.session_state:
    st.session_state.embedding_model = None  # Placeholder for the embedding model

if "llm_type" not in st.session_state:
    st.session_state.llm_type = LOCAL_LLM

if "llm_name" not in st.session_state:
    st.session_state.llm_name = DEFAULT_LOCAL_LLM

if "llm_model" not in st.session_state:
    st.session_state.llm_model = None

if "client" not in st.session_state:
    st.session_state.client = chromadb.PersistentClient("db")

if "collection" not in st.session_state:
    st.session_state.collection = None

if "search_option" not in st.session_state:
    st.session_state.search_option = "Vector Search"

if "open_dialog" not in st.session_state:
    st.session_state.open_dialog = None

if "source_data" not in st.session_state:
    st.session_state.source_data = "UPLOAD"

if "chunks_df" not in st.session_state:
    st.session_state.chunks_df = pd.DataFrame()

if "random_collection_name" not in st.session_state:
    st.session_state.random_collection_name = None

if "chunk_size" not in st.session_state:
    st.session_state.chunk_size = 200

if "chunk_overlap" not in st.session_state:
    st.session_state.chunk_overlap = 10

# --- End of initialization


if "graph_query" not in st.session_state:
    st.session_state.graph_query = """MATCH (n)-[r]->(m) RETURN n, r, m LIMIT 100"""

# Sidebar settings
st.sidebar.header("Settings")

# Chunk size input
st.session_state.chunk_size = st.sidebar.number_input(
    "Chunk Size",
    min_value=10, 
    max_value=1000, 
    value=200, 
    step=10, 
    help="Set the size of each chunk in terms of tokens.",

)
st.session_state.chunk_overlap = st.sidebar.number_input(
    "Chunk Overlap",
    min_value=0, 
    max_value=1000, 
    step=10, 
    value=st.session_state.chunk_size // 10,
    help="Set the size of each chunk in terms of tokens.",

)


st.session_state.number_docs_retrieval = st.sidebar.number_input(
    "Number of documnents retrieval", 
    min_value=1, 
    max_value=50,
    value=3,
    step=1,
    help="Set the number of document which will be retrieved."
)

header_i = 1
# Language selection popup
header_text = "{}. Setup Language".format(header_i)
st.header(header_text)
language_choice = st.radio(
    "Select language:", [
        NONE,
        ENGLISH, 
        VIETNAMESE
    ],
    index=0
    )

# Switch embedding model based on language choice
if language_choice == ENGLISH:
    if st.session_state.get("language") != EN:
        st.session_state.language = EN
        # Only load the model if it hasn't been loaded before
        if st.session_state.get("embedding_model_name") != 'all-MiniLM-L6-v2':
            st.session_state.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
            st.session_state.embedding_model_name = 'all-MiniLM-L6-v2'
        st.success("Using English embedding model: all-MiniLM-L6-v2")

elif language_choice == VIETNAMESE:
    if st.session_state.get("language") != VI:
        st.session_state.language = VI
        # Only load the model if it hasn't been loaded before
        if st.session_state.get("embedding_model_name") != 'keepitreal/vietnamese-sbert':
            st.session_state.embedding_model = SentenceTransformer('keepitreal/vietnamese-sbert')
            st.session_state.embedding_model_name = 'keepitreal/vietnamese-sbert'
        st.success("Using Vietnamese embedding model: keepitreal/vietnamese-sbert")



# Step 1: File Upload (CSV, JSON, PDF, or DOCX) and Column Detection

header_i += 1
st.header(f"{header_i}. Setup data source")
st.subheader(f"{header_i}.1. Upload data ", divider=True)
uploaded_files = st.file_uploader(
    "Upload CSV, JSON, PDF, or DOCX files", 
    # type=["csv", "json", "pdf", "docx"], 
    accept_multiple_files=True
)

# Initialize a variable for tracking the success of saving the data
st.session_state.data_saved_success = False

# Ensure `df` is only accessed if it's been created and has columns
if uploaded_files is not None:
    all_data = []
    
    for uploaded_file in uploaded_files:
        print(uploaded_file.type)
        # Determine file type and read accordingly
        if uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
            all_data.append(df)

        elif uploaded_file.name.endswith(".json"):
            json_data = json.load(uploaded_file)
            df = pd.json_normalize(json_data)  # Normalize JSON to a flat DataFrame format
            all_data.append(df)

        elif uploaded_file.name.endswith(".pdf"):
            # Extract text from PDF
            pdf_text = []
            with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
                for page in pdf.pages:
                    pdf_text.append(page.extract_text())

            # Convert PDF text into a DataFrame (assuming one column for simplicity)
            df = pd.DataFrame({"content": pdf_text})
            all_data.append(df)

        elif uploaded_file.name.endswith(".docx") or uploaded_file.name.endswith(".doc"):
            # Extract text from DOCX
            doc = Document(io.BytesIO(uploaded_file.read()))
            docx_text = [para.text for para in doc.paragraphs if para.text]

            # Convert DOCX text into a DataFrame (assuming one column for simplicity)
            df = pd.DataFrame({"content": docx_text})
            all_data.append(df)
        elif uploaded_file.name.endswith(".xlsx") or uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            # Read Excel file
            try:
                df = pd.read_excel(
                    uploaded_file,
                    engine="openpyxl"
                    )
                all_data.append(df)
            except Exception as e:
                st.error(f"Error reading file: {str(e)}")
        else:
            st.error("Unsupported file format.")


    # Concatenate all data into a single DataFrame
    if all_data:
        df = pd.concat(all_data, ignore_index=True)
        st.dataframe(df)

        doc_ids = [str(uuid.uuid4()) for _ in range(len(df))]

        if "doc_ids" not in st.session_state:
            st.session_state.doc_ids = doc_ids

        # Add or replace the '_id' column in the DataFrame
        df['doc_id'] = st.session_state.doc_ids

        st.subheader("Chunking")

        # **Ensure `df` is not empty before calling selectbox**
        if not df.empty:
            # Display selectbox to choose the column for vector search
            index_column = st.selectbox("Choose the column to index (for vector search):", df.columns)
            st.write(f"Selected column for indexing: {index_column}")
        else:
            st.warning("The DataFrame is empty, please upload valid data.")
            
        # Disable the "AgenticChunker" option if the API key is not provided
        chunk_options = [
            NO_CHUNKING,
            "RecursiveTokenChunker", 
            "SemanticChunker",
            "AgenticChunker",
        ]

        # Step 4: Chunking options
        if st.session_state.get("llm_choice") == llm_options["Online"] and not st.session_state.get("llm_api_key") and st.session_state.get("chunkOption") == "AgenticChunker":
            currentChunkerIdx = 0
            st.session_state.chunkOption = NO_CHUNKING
            notify("You have to setup the GEMINI API KEY FIRST in the Setup LLM Section for using Online Agentic Chunker", "error")
        elif not st.session_state.get("chunkOption"):
            currentChunkerIdx = 0
            st.session_state.chunkOption = NO_CHUNKING
        else:
            currentChunkerIdx = chunk_options.index(st.session_state.get("chunkOption")) 
        
        st.radio(
            "Please select one of the options below.",
            chunk_options,
            captions=[
                "Keep the original document",
                "Recursively chunks text into smaller, meaningful token groups based on specific rules or criteria.",
                "Chunking with semantic comparison between chunks",
                "Let LLM decide chunking (requires Gemini API)"
            ],
            key="chunkOption",
            index=currentChunkerIdx
        )
        
        chunkOption = st.session_state.get("chunkOption")
        
        if chunkOption == "SemanticChunker":
            embedding_option = st.selectbox(
                "Choose the embedding method for Semantic Chunker:",
                ["TF-IDF", "Sentence-Transformers"]
            )
        chunk_records = []

        print('--going here', chunkOption)


        # Iterate over rows in the original DataFrame
        for index, row in df.iterrows():

            # For "No Chunking" option, treat the selected index column as a single "chunk"
            chunker = None
            selected_column_value = row[index_column]
            chunks = []
            if not(type(selected_column_value) == str and len(selected_column_value) > 0):
                continue
            if chunkOption == NO_CHUNKING:
                # Use the selected index_column
                chunks = [selected_column_value]
                
            # For "RecursiveTokenChunker" option, split text from the selected index column into smaller chunks
            elif chunkOption == "RecursiveTokenChunker":
                chunker = RecursiveTokenChunker(
                    chunk_size=st.session_state.chunk_size,
                    chunk_overlap=st.session_state.chunk_overlap

                )
                chunks = chunker.split_text(selected_column_value)
                
            elif chunkOption == "SemanticChunker":
                if embedding_option == "TF-IDF":
                    chunker = ProtonxSemanticChunker(
                        embedding_type="tfidf",
                    )
                else:
                    chunker = ProtonxSemanticChunker(
                        embedding_type="transformers", 
                        model="all-MiniLM-L6-v2",
                    )
                chunks = chunker.split_text(selected_column_value)
            elif chunkOption == "AgenticChunker" and st.session_state.get("llm_choice") == llm_options["Online"]  and st.session_state.get("llm_api_key"):
                chunker = LLMAgenticChunkerv2(st.session_state.get("llm_model"))
                chunks = chunker.split_text(selected_column_value)
            elif chunkOption == "AgenticChunker" and st.session_state.get("llm_choice") == llm_options["Local (Ollama)"] and st.session_state.get("local_llms"):
                chunker = LLMAgenticChunkerv2(st.session_state.get("local_llms"))
                chunks = chunker.split_text(selected_column_value)
            # For each chunk, add a dictionary with the chunk and original_id to the list
            for chunk in chunks:
                chunk_record = {**row.to_dict(), 'chunk': chunk}
                
                # Rearrange the dictionary to ensure 'chunk' and '_id' come first
                chunk_record = {
                    'chunk': chunk_record['chunk'],
                    # '_id': str(uuid.uuid4()),
                    **{k: v for k, v in chunk_record.items() if k not in ['chunk', '_id']}
                }
                chunk_records.append(chunk_record)

        # Convert the list of dictionaries to a DataFrame
        st.session_state.chunks_df = pd.DataFrame(chunk_records)



if "chunks_df" in st.session_state and len(st.session_state.chunks_df) > 0:
    # Display the result
    st.write("Number of chunks:", len(st.session_state.chunks_df))
    st.dataframe(st.session_state.chunks_df)

# Button to save data
if st.button("Save Data"):
    try:
        # Check if the collection exists, if not, create a new one
        if st.session_state.collection is None:
            if uploaded_files:
                first_file_name = os.path.splitext(uploaded_files[0].name)[0]  # Get file name without extension
                collection_name = f"rag_collection_{clean_collection_name(first_file_name)}"
            else:
                # If no file name is available, generate a random collection name
                collection_name = f"rag_collection_{uuid.uuid4().hex[:8]}"
        
            st.session_state.random_collection_name = collection_name
            st.session_state.collection = st.session_state.client.get_or_create_collection(
                name=st.session_state.random_collection_name,
                metadata={"description": "A collection for RAG system"},
            )

        # Define the batch size
        batch_size = 256

        # Split the DataFrame into smaller batches
        df_batches = divide_dataframe(st.session_state.chunks_df, batch_size)

        # Check if the dataframe has data, otherwise show a warning and skip the processing
        if not df_batches:
            st.warning("No data available to process.")
        else:
            num_batches = len(df_batches)

            # Initialize progress bar
            progress_text = "Saving data to Chroma. Please wait..."
            my_bar = st.progress(0, text=progress_text)

            # Process each batch
            for i, batch_df in enumerate(df_batches):
                if batch_df.empty:
                    continue  # Skip empty batches (just in case)
                
                process_batch(batch_df, st.session_state.embedding_model, st.session_state.collection)

                # Update progress dynamically for each batch
                progress_percentage = int(((i + 1) / num_batches) * 100)
                my_bar.progress(progress_percentage, text=f"Processing batch {i + 1}/{num_batches}")

                time.sleep(0.1)  # Optional sleep to simulate processing time

            # Empty the progress bar once completed
            my_bar.empty()

            st.success("Data saved to Chroma vector store successfully!")
            st.markdown("Collection name: `{}`".format(st.session_state.random_collection_name))
            st.session_state.data_saved_success = True

    except Exception as e:
        st.error(f"Error saving data to Chroma: {str(e)}")

# Set up the interface
st.subheader(f"{header_i}.2. Or load from saved collection", divider=True)
if st.button("Load from saved collection"):
    st.session_state.open_dialog = "LIST_COLLECTION"
    def load_func(collection_name):
        st.session_state.collection = st.session_state.client.get_collection(
            name=collection_name
        )
        st.session_state.random_collection_name = collection_name
        st.session_state.data_saved_success = True
        st.session_state.source_data = DB
        data = st.session_state.collection.get(
            include=[
                "documents", 
                "metadatas"
            ],
        )
        metadatas = data["metadatas"]
        column_names = []
        if len(metadatas) > 0 and len(metadatas[0].keys()) > 0:
            column_names.extend(metadatas[0].keys())
            column_names = list(set(column_names))

        st.session_state.chunks_df = pd.DataFrame(metadatas, columns=column_names)

    def delete_func(collection_name):
        st.session_state.client.delete_collection(name=collection_name)
    
    list_collection(st.session_state, load_func, delete_func)
        

header_i += 1
header_text = "{}. Setup data ✅".format(header_i) if st.session_state.data_saved_success else "{}. Setup data".format(header_i)
st.header(header_text)

if st.session_state.data_saved_success:
    st.markdown("✅ **Data Saved Successfully!**")




# Step 3: Define which columns LLMs should answer from
if "random_collection_name" in st.session_state and st.session_state.random_collection_name is not None and st.session_state.chunks_df is not None:
    st.session_state.columns_to_answer = st.multiselect(
        "Select one or more columns LLMs should answer from (multiple selections allowed):", 
        st.session_state.chunks_df.columns
    )

# Step 2: Setup LLMs (Gemini Only)
header_i += 1
header_text_llm = "{}. Setup LLMs".format(header_i)
st.header(header_text_llm)
# Example user selection
st.selectbox(
    "Choose Model Source:", 
    ["Online", "Local (Ollama)"],
    index=list(llm_options.values()).index(llm_options["Local (Ollama)"]),
    key="llm_choice"
)
if st.session_state.get("llm_choice") == llm_options["Online"]:
    # Initialize a variable for tracking if the API key was entered successfully
    api_key_success = False
    st.session_state.llm_type = ONLINE_LLM

    if st.session_state.llm_type == ONLINE_LLM:
        st.session_state.llm_name = st.selectbox(
            "Select the LLM model to run", 
            options=[GEMINI, OPENAI], 
            index=0
        )
   
    # Input API key
    if st.session_state.llm_name == GEMINI:
        st.markdown("Obtain the API key from the [Google AI Studio](https://aistudio.google.com/app/apikey).")
    elif st.session_state.llm_name == OPENAI:
        st.markdown("Obtain the API key from the [OpenAI API page](https://platform.openai.com/account/api-keys).")
    
    st.text_input(
        "Enter your API Key:", 
        type="password", 
        key="llm_api_key",
        value=os.getenv("GEMINI_API_KEY") if st.session_state.llm_name == GEMINI else os.getenv("OPENAI_API_KEY"),
    ) 
    if st.session_state.get('llm_api_key'):
        st.success("API Key saved successfully!")
        if st.session_state.llm_name == GEMINI:
            st.session_state.llm_model = OnlineLLMs(
                name=GEMINI,
                api_key=st.session_state.get('llm_api_key'),
                model_version="gemini-1.5-pro"
                )
        elif st.session_state.llm_name == OPENAI:
            st.session_state.llm_model = OnlineLLMs(
                name=OPENAI,
                api_key=st.session_state.get('llm_api_key'),
                model_version="gpt-4o",
                )
        else:
            st.warning("Please select a model to run.")

        api_key_success = True

    # Show blue tick if API key was entered successfully
    if api_key_success:
        st.markdown("✅ **API Key Saved Successfully!**")
elif st.session_state.get("llm_choice") == llm_options["Local (Ollama)"]:
    st.markdown("Please install and run [Docker](https://docs.docker.com/engine/install/) before running Ollama locally.")
    # Install and run Ollama Docker container based on hardware
    if st.button("Initialize Ollama Container"):
        with st.spinner("Setting up Ollama container..."):
            run_ollama_container()
        
    
    model_format = st.radio(
        label="### Select the model format",
        options=["Normal", "High Performance"],
        captions=["HuggingFace normal format", "HuggingFace GGUF format"],
        index=0,
    )

    if model_format == "Normal":
        selected_model = st.selectbox("Select a model to run", list(OLLAMA_MODEL_OPTIONS.keys()))
        real_name_model = OLLAMA_MODEL_OPTIONS[selected_model]
    elif model_format == "High Performance":
        selected_model = st.selectbox("Select a model to run", list(GGUF_MODEL_OPTIONS.keys()))
        real_name_model = GGUF_MODEL_OPTIONS[selected_model]

        st.markdown("Or type the name of the model to confirm. The list of models from HuggingFace in GGUF format [here](https://huggingface.co/models?library=gguf&sort=trending)")
        
        type_selected_model = st.text_input(
            "The GGUF model's name should be: `hf.co/{username}/{repository}`. Check the [document](https://huggingface.co/docs/hub/en/ollama)", 
            key="type_selected_model")
        if type_selected_model:
            real_name_model = type_selected_model

    if st.button("Run Selected Model"):
        localLLms = run_ollama_model(real_name_model)
        st.session_state.llm_name = real_name_model
        st.session_state.llm_type = LOCAL_LLM
        st.session_state.local_llms = localLLms


st.sidebar.subheader("All configurations:")
st.sidebar.markdown(f"1. Collection name: **{st.session_state.collection.name if st.session_state.collection else 'No collection'}**")
st.sidebar.markdown(f"2. LLM model: **{st.session_state.llm_name if 'llm_name' in st.session_state else 'Not selected'}**")
st.sidebar.markdown(f"3. Local or APIs: **{st.session_state.llm_type if 'llm_type' in st.session_state else 'Not specified'}**")
st.sidebar.markdown(f"4. Language: **{st.session_state.language}**")
st.sidebar.markdown(f"5. Embedding Model: **{st.session_state.embedding_model.__class__.__name__ if st.session_state.embedding_model else 'None'}**")
st.sidebar.markdown(f"6. Chunk Size: **{st.session_state.chunk_size}**")
st.sidebar.markdown(f"7. Number of Documents Retrieval: **{st.session_state.number_docs_retrieval}**")
st.sidebar.markdown(f"8. Data Saved: **{'Yes' if st.session_state.data_saved_success else 'No'}**")
st.sidebar.markdown(f"9. LLM API Key Set: **{'Yes' if st.session_state.get('llm_api_key') else 'No'}**")
if st.session_state.get('chunkOption'):
    st.sidebar.markdown(f"10. Chunking Option: **{st.session_state.chunkOption}**")


header_i += 1
header_text = "{}. [Experimental] Graph RAG".format(header_i)
st.header(header_text)

# Display graph

def update_graph_query():
    st.session_state.graph_query = st.session_state.temp_graph_query

if st.button("Extract Graph"):
    if not st.session_state.get("llm_api_key"):
        notify("You have to setup the API KEY FIRST in the Setup LLM Section", "error")
    else:
        if st.session_state.llm_type == ONLINE_LLM:
            if st.session_state.llm_name == GEMINI:
                llm_model_for_graph_rag = ChatGoogleGenerativeAI(
                    model="gemini-1.5-pro",
                    google_api_key=st.session_state.get("llm_api_key")
                )
            elif st.session_state.llm_name == OPENAI:

                llm_model_for_graph_rag = ChatOpenAI(
                    model="gpt-4o",
                    openai_api_key=st.session_state.get("llm_api_key")
                )
                
            st.session_state.graph_rag = GraphRAG(llm_model_for_graph_rag)

            langchain_documents = [
                langchainDocument(
                    page_content=
                    (
                        f"{row['chunk']}"
                    )
                ) 
                for _, row in st.session_state.chunks_df.iterrows()
            ]
            
            st.session_state.graph_rag.create_graph(langchain_documents)

            # e_r = st.session_state.graph_rag.extract_entities_and_relationships()
        else:
            notify("Only support online model for now.", "error")

if "graph_rag" in st.session_state and st.session_state.graph_rag is not None:
    graph_query = st.text_area(
        "Graph Query",
        st.session_state.graph_query,
        key="temp_graph_query",
        on_change=update_graph_query
    )

    if st.button("Visualize Graph"):
        data = st.session_state.graph_rag.query_graph(graph_query)
        st.session_state.graph_rag.visualize_graph(data)

header_i += 1
header_text_llm = "{}. Set up search algorithms".format(header_i)
st.header(header_text_llm)

st.radio(
    "Please select one of the options below.",
    [
        # "Keywords Search", 
        "Vector Search", 
        "Hyde Search"],
    captions = [
        # "Search using traditional keyword matching",
        "Search using vector similarity",
        "Search using the HYDE algorithm"
    ],
    key="search_option",
    index=0,
)


header_i += 1
st.header("{}. Export Chatbot".format(header_i))
if st.button("Export Chatbot"):
    st.write("Exporting chatbot...")

    # Save session state to a JSON file
    def save_session_state(file_path="pages/session_state.json"):
        # Ensure the directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        # Required fields to export
        required_fields = [
            "random_collection_name", 
            "number_docs_retrieval", 
            "embedding_model_name", 
            "llm_type", 
            "llm_name",
            "columns_to_answer",
            "search_option"
        ]

        # Check if all required fields are present in session state
        missing_fields = [field for field in required_fields if field not in st.session_state]
        if missing_fields:
            st.error(f"Missing required fields: {', '.join(missing_fields)}")
            return

        # Check if llm_type is 'local_llm'
        if st.session_state["llm_type"] != "local_llm":
            st.error("Only support exporting llms local.")
            return
        
        # Filter session state to only include specified fields and serializable types
        session_data = {
            key: value for key, value in st.session_state.items() 
            if key in required_fields and isinstance(value, (str, int, float, bool, list, dict))
        }

        # Save to JSON file
        with open(file_path, "w") as file:
            json.dump(session_data, file)
        
        st.success("Chatbot exported successfully!")

    save_session_state()

# Step 4: Interactive Chatbot
header_i += 1
st.header("{}. Interactive Chatbot".format(header_i))

# Initialize chat history in session state
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# URL of the Flask API

# Display the chat history using chat UI
for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])


if prompt := st.chat_input("What is up?"):
    # Add user message to chat history
    st.session_state.chat_history.append({"role": USER, "content": prompt})
    # Display user message in chat message container
    with st.chat_message(USER):
        st.markdown(prompt)
    # Display assistant response in chat message container
    # Prepare the payload for the request

    with st.chat_message(ASSISTANT):
        if st.session_state.collection is not None:
            # Combine retrieved data to enhance the prompt based on selected columns
            metadatas, retrieved_data = [], ""
            if st.session_state.columns_to_answer:
                if st.session_state.search_option == "Vector Search":
                    metadatas, retrieved_data = vector_search(
                        st.session_state.embedding_model, 
                        prompt, 
                        st.session_state.collection, 
                        st.session_state.columns_to_answer,
                        st.session_state.number_docs_retrieval
                    )
                    
                    enhanced_prompt = """The prompt of the user is: "{}". Answer it based on the following retrieved data: \n{}""".format(prompt, retrieved_data)

                elif st.session_state.search_option == "Keywords Search":
                    metadatas, retrieved_data = keywords_search(
                        prompt,
                        st.session_state.collection,
                        st.session_state.columns_to_answer,
                        st.session_state.number_docs_retrieval
                    )

                    enhanced_prompt = """The prompt of the user is: "{}". Answer it based on the following retrieved data: \n{}""".format(prompt, retrieved_data)

                elif st.session_state.search_option == "Hyde Search":
              
                    if st.session_state.llm_type == ONLINE_LLM:
                        model = st.session_state.llm_model
                    else:
                        model = st.session_state.local_llms


                    metadatas, retrieved_data = hyde_search(
                        model,
                        st.session_state.embedding_model,
                        prompt,
                        st.session_state.collection,
                        st.session_state.columns_to_answer,
                        st.session_state.number_docs_retrieval,
                        num_samples=1
                    )

                    enhanced_prompt = """The prompt of the user is: "{}". Answer it based on the following retrieved data: \n{}""".format(prompt, retrieved_data)

                
                if metadatas:
                    flattened_metadatas = [item for sublist in metadatas for item in sublist]  # Flatten the list of lists
                    
                    # Convert the flattened list of dictionaries to a DataFrame
                    metadata_df = pd.DataFrame(flattened_metadatas)
                    
                    # Display the DataFrame in the sidebar
                 
                    st.sidebar.subheader("Retrieval data")
                    st.sidebar.dataframe(metadata_df)
                    st.sidebar.subheader("Full prompt for LLM")
                    st.sidebar.markdown(enhanced_prompt)
                else:
                    st.sidebar.write("No metadata to display.")

                if st.session_state.llm_type == ONLINE_LLM:
                    # Generate content using the selected LLM model
                    if "llm_model" in st.session_state and st.session_state.llm_model is not None:
                        response = st.session_state.llm_model.generate_content(enhanced_prompt)

                        # Display the extracted content in the Streamlit app
                        st.markdown(response)

                        # Update chat history
                        st.session_state.chat_history.append({"role": ASSISTANT, "content": response})
                    else:
                        st.warning("Please select a model to run.")
                elif st.session_state.llm_type == LOCAL_LLM:
                    if "local_llms" in st.session_state and st.session_state.local_llms is not None:
                        response = st.session_state.local_llms.generate_content(enhanced_prompt)
                        st.markdown(response)
                        st.session_state.chat_history.append({"role": ASSISTANT, "content": response})
                    else:
                        st.warning("Please select a model to run.")
            else:
                st.warning("Please select columns for the chatbot to answer from.")
        else:
            st.error("No collection found. Please upload data and save it first.")



    