import streamlit as st
import pandas as pd
import json
import pdfplumber
import io
from docx import Document
import re
from openai import OpenAI
from graph_rag import GraphRAG
from langchain_openai import ChatOpenAI
from langchain.schema import Document as langchainDocument

st.logo("https://storage.googleapis.com/mle-courses-prod/users/61b6fa1ba83a7e37c8309756/private-files/678dadd0-603b-11ef-b0a7-998b84b38d43-ProtonX_logo_horizontally__1_.png")

DEFAULT_GRAPH_QUERY = "MATCH (n)-[r]->(m) RETURN n, r, m LIMIT 100"

# Initialize session state variables
if "graph_query" not in st.session_state:
    st.session_state.graph_query = DEFAULT_GRAPH_QUERY

if "temp_graph_query" not in st.session_state:
    st.session_state.temp_graph_query = DEFAULT_GRAPH_QUERY

if "graph_rag" not in st.session_state:
    st.session_state.graph_rag = None

if "processing_successful" not in st.session_state:
    st.session_state.processing_successful = False

def update_graph_query():
    st.session_state.graph_query = st.session_state.temp_graph_query

# Initialize the page
st.header("Drag and Drop Data Loader")
st.markdown("Upload your data files (CSV, JSON, PDF, DOCX, Excel) to load and preview them.")

# Step 1: API Key Input
st.subheader("0. Enter OpenAI API Key", divider=True)
api_key = st.text_input(
    "Enter your OpenAI API Key",
    type="password",
    placeholder="sk-..."
)

if not api_key:
    st.warning("Please enter your OpenAI API Key to proceed.")

# Step 2: File Upload (CSV, JSON, PDF, DOCX, Excel)
st.subheader("1. Upload Data Files", divider=True)
uploaded_files = st.file_uploader(
    "Upload CSV, JSON, PDF, DOCX, or Excel files",
    accept_multiple_files=True,
    key="file_uploader_unique"
)

if uploaded_files and api_key:
    all_data = []
    for uploaded_file in uploaded_files:
        if uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
            all_data.append(df)
        elif uploaded_file.name.endswith(".json"):
            json_data = json.load(uploaded_file)
            df = pd.json_normalize(json_data)
            all_data.append(df)
        elif uploaded_file.name.endswith(".pdf"):
            pdf_text = []
            with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
                for page in pdf.pages:
                    pdf_text.append(page.extract_text())
            df = pd.DataFrame({"content": pdf_text})
            all_data.append(df)
        elif uploaded_file.name.endswith(".docx") or uploaded_file.name.endswith(".doc"):
            doc = Document(io.BytesIO(uploaded_file.read()))
            docx_text = [para.text for para in doc.paragraphs if para.text]
            df = pd.DataFrame({"content": docx_text})
            all_data.append(df)
        elif uploaded_file.name.endswith(".xlsx") or uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            try:
                df = pd.read_excel(uploaded_file, engine="openpyxl")
                all_data.append(df)
            except Exception as e:
                st.error(f"Error reading file: {str(e)}")
        else:
            st.error("Unsupported file format.")
    
    # Concatenate all data into a single DataFrame
    if all_data:
        combined_df = pd.concat(all_data, ignore_index=True)
        st.write("### Preview of Uploaded Data:")
        st.dataframe(combined_df)
        
        # Column Selection for RAG Conversion
        selected_column = st.selectbox(
            "Select a column to convert to RAG:",
            combined_df.columns
        )
        
        if st.button("Process Column for RAG Conversion"):
            text_data = combined_df[selected_column].dropna().astype(str).str.cat(sep=' ')
            st.write("### Selected Column Data for RAG Conversion")

            # Run Extraction
            with st.spinner("Processing text with OpenAI..."):
                try:
                    llm_model_for_graph_rag = ChatOpenAI(
                        model="gpt-4o",
                        openai_api_key=api_key
                    )
                    st.session_state.graph_rag = GraphRAG(llm_model_for_graph_rag)

                    langchain_documents = [
                        langchainDocument(
                            page_content=str(row)
                        )
                        for row in combined_df[selected_column].dropna()
                    ]

                    st.session_state.graph_rag.create_graph(langchain_documents)
                    st.session_state.processing_successful = True  # Mark as successful
                    st.success("Data successfully processed! You can now visualize the graph.")
                except Exception as e:
                    st.session_state.processing_successful = False
                    st.error(f"Error during entity extraction: {e}")
    else:
        st.warning("No valid data to display.")

# Display Visualization Button only if processing was successful
if st.session_state.processing_successful:
    st.subheader("2. Visualize Graph", divider=True)
    if st.button("Visualize Graph"):
        st.session_state.graph_query = st.text_area(
            "Graph Query",
            st.session_state.graph_query,
            key="temp_graph_query",
            on_change=update_graph_query
        )
        if st.session_state.graph_query.strip() == "":
            st.warning("Graph query cannot be empty. Please enter a valid query.")
        else:
            try:
                data = st.session_state.graph_rag.query_graph(st.session_state.graph_query)
                graph_output = st.session_state.graph_rag.visualize_graph(data)
                
                # Check visualization type
                if isinstance(graph_output, str):
                    st.components.v1.html(graph_output, height=600)
                elif hasattr(graph_output, 'to_image'):
                    st.image(graph_output.to_image(format='png'))
                elif hasattr(graph_output, 'show'):
                    st.plotly_chart(graph_output)
                else:
                    st.write(graph_output)
            except Exception as e:
                st.error(f"Error visualizing graph: {e}")
